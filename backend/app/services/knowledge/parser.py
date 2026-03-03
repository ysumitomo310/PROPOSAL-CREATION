"""ナレッジパーサー（TASK-B02: BPDParser, TASK-B03: ModuleOverviewParser）

BPDドキュメント（docx 3ファイルセット）と Discovery WS PDF を解析し、
Neo4j投入用の構造化データを生成する。
"""

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path

import pdfplumber
from docx import Document
from langchain_core.messages import HumanMessage, SystemMessage

from app.core.llm_client import LLMClient
from app.services.knowledge.master import ModuleClassificationService

logger = logging.getLogger(__name__)


# =============================================================================
# Data classes
# =============================================================================


@dataclass
class ScopeItemData:
    """BPDから抽出されたScopeItem構造化データ。"""

    id: str  # "SAP-1B4"
    product: str  # "SAP S/4 HANA Public Edition"
    product_namespace: str  # "SAP"
    module: str  # "SD"
    scope_item_id: str  # "1B4"
    function_name: str  # "受注から入金"
    description: str  # LLM生成の機能要約（JA）
    description_en: str  # EN版テキスト
    business_domain: str  # "販売"
    capability_level: str = "standard"
    keywords: list[str] = field(default_factory=list)
    source_doc: str = ""
    product_version: str = "SAP S/4HANA Cloud Public Edition 2602"
    relations: dict[str, list[str]] = field(default_factory=dict)


@dataclass
class ModuleOverviewData:
    """Discovery WS PDFから抽出されたModuleOverview構造化データ。"""

    id: str  # "MO-SD-inventory-sales"
    product: str
    product_namespace: str  # "SAP"
    module: str  # "SD"
    module_name: str  # "SD在庫販売ソリューション"
    summary: str  # LLM生成
    source_doc: str
    page_count: int
    covers_scope_items: list[str] = field(default_factory=list)


# =============================================================================
# BPD Parser (TASK-B02)
# =============================================================================

# Heading 1 セクション名（B02a Spikeで確認済み）
_JA_SECTION_PURPOSE = "目的"
_JA_SECTION_PREREQUISITES = "前提条件"
_JA_SECTION_OVERVIEW = "概要テーブル"
_JA_SECTION_PROCEDURES = "テスト手順"
_JA_SECTION_APPENDIX = "付録"

# Scope Item ID参照パターン: "(2EL)", "(BJE)" 形式、またはテーブル内 "1XN - 説明"
_SI_REF_PAREN = re.compile(r'\(([A-Z0-9]{2,5})\)')
_SI_REF_TABLE = re.compile(r'^([A-Z0-9]{2,5})\s*[-–—]')


@dataclass
class JASections:
    """JA版BPDから抽出されたセクション。"""

    purpose: str = ""
    prerequisites: str = ""
    business_conditions: list[tuple[str, str]] = field(default_factory=list)  # [(si_id, condition)]
    overview_table: list[dict] = field(default_factory=list)  # [{process, role, app, result}]
    procedures: list[str] = field(default_factory=list)  # H2テスト手順名のリスト


class BPDParser:
    """BPD 3ファイルセットからScopeItemDataを生成。"""

    def __init__(
        self,
        llm_client: LLMClient | None = None,
        master_service: ModuleClassificationService | None = None,
    ) -> None:
        self._llm = llm_client
        self._master = master_service

    async def parse_scope_item(
        self,
        ja_docx_path: Path,
        en_docx_path: Path | None = None,
        xlsx_path: Path | None = None,
    ) -> ScopeItemData:
        """3ファイルセットからScopeItem構造化データを生成。"""
        # 1. プレフィクス抽出
        scope_item_id = self._extract_prefix(ja_docx_path.name)

        # 2. JA docx 解析
        ja_sections = self._parse_ja_docx(ja_docx_path)
        function_name = self._extract_function_name(ja_docx_path)

        # 3. EN docx 解析
        en_text = ""
        if en_docx_path and en_docx_path.exists():
            en_text = self._extract_en_purpose(en_docx_path)

        # 4. モジュール情報取得
        module = "OTHER"
        module_name_ja = "その他"
        business_domain = "その他"
        if self._master:
            info = await self._master.get_module(scope_item_id)
            if info:
                module = info.module
                module_name_ja = info.module_name_ja
                business_domain = info.business_domain

        # 5. リレーション抽出
        relations = self._extract_relations(
            ja_sections.business_conditions, ja_sections.procedures
        )

        # 6. description 生成 (LLM)
        description = ja_sections.purpose
        if self._llm and ja_sections.purpose:
            description = await self._generate_description(
                purpose=ja_sections.purpose,
                procedures=ja_sections.procedures,
                module_context=module_name_ja,
            )

        # 7. キーワード抽出
        keywords = self._extract_keywords(function_name, ja_sections)

        return ScopeItemData(
            id=f"SAP-{scope_item_id}",
            product="SAP S/4 HANA Public Edition",
            product_namespace="SAP",
            module=module,
            scope_item_id=scope_item_id,
            function_name=function_name,
            description=description,
            description_en=en_text,
            business_domain=business_domain,
            keywords=keywords,
            source_doc=ja_docx_path.name,
            relations=relations,
        )

    def _extract_prefix(self, filename: str) -> str:
        """ファイル名からScope Item IDプレフィクスを抽出。"""
        # "16R_S4CLD2602_BPD_JA_JP.docx" → "16R"
        return filename.split("_S4CLD")[0]

    def _extract_function_name(self, docx_path: Path) -> str:
        """タイトルテーブル（Table 0, Row 1）からfunction_nameを抽出。"""
        doc = Document(str(docx_path))
        if doc.tables and len(doc.tables[0].rows) >= 2:
            title = doc.tables[0].rows[1].cells[-1].text.strip()
            # "(XXX_JP)" サフィックスを除去
            title = re.sub(r'\s*\([A-Z0-9]+_JP\)\s*$', '', title)
            return title
        return docx_path.stem

    def _parse_ja_docx(self, docx_path: Path) -> JASections:
        """JA版docxのセクション構造を解析。"""
        doc = Document(str(docx_path))
        sections = JASections()
        current_h1 = ""
        current_h2 = ""
        buffer = []

        for para in doc.paragraphs:
            style = para.style.name
            text = para.text.strip()

            if style == "Heading 1":
                # 前のセクションを保存
                self._flush_section(sections, current_h1, "\n".join(buffer))
                current_h1 = text
                current_h2 = ""
                buffer = []
            elif style == "Heading 2":
                if current_h1 == _JA_SECTION_PROCEDURES:
                    # テスト手順配下のH2 → プロセスステップ名
                    sections.procedures.append(text)
                current_h2 = text
            elif text:
                buffer.append(text)

        self._flush_section(sections, current_h1, "\n".join(buffer))

        # テーブルからビジネス条件とOverviewを抽出
        for table in doc.tables:
            self._parse_table(table, sections)

        return sections

    def _flush_section(self, sections: JASections, h1: str, text: str) -> None:
        if _JA_SECTION_PURPOSE in h1:
            sections.purpose = text
        elif _JA_SECTION_PREREQUISITES in h1:
            sections.prerequisites = text

    def _parse_table(self, table, sections: JASections) -> None:
        """テーブルを解析してビジネス条件やOverviewを抽出。"""
        if not table.rows:
            return
        headers = [cell.text.strip() for cell in table.rows[0].cells]

        # ビジネス条件テーブル: ["スコープアイテム", "ビジネス条件"]
        if len(headers) >= 2 and "スコープアイテム" in headers[0]:
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 2 and cells[0]:
                    sections.business_conditions.append((cells[0], cells[1]))

        # 概要テーブル: ["プロセスステップ", "ビジネスロール", "トランザクション/アプリ", "予想される結果"]
        elif len(headers) >= 4 and "プロセスステップ" in headers[0]:
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 4 and cells[0]:
                    sections.overview_table.append({
                        "process": cells[0],
                        "role": cells[1],
                        "app": cells[2],
                        "result": cells[3],
                    })

    def _extract_relations(
        self,
        business_conditions: list[tuple[str, str]],
        procedures: list[str],
    ) -> dict[str, list[str]]:
        """ビジネス条件とプロセス手順からScope Item ID参照を抽出。"""
        related_ids = set()

        # ビジネス条件テーブルから
        for si_text, _ in business_conditions:
            match = _SI_REF_TABLE.match(si_text)
            if match:
                related_ids.add(match.group(1))

        # プロセス手順のヘッダーから "(2EL)" パターン
        for proc in procedures:
            for match in _SI_REF_PAREN.finditer(proc):
                related_ids.add(match.group(1))

        return {"related": sorted(related_ids)} if related_ids else {}

    def _extract_keywords(self, function_name: str, sections: JASections) -> list[str]:
        """function_name とプロセス手順からキーワードを抽出。"""
        keywords = set()
        # function_name を分割
        for part in re.split(r'[、/・()（）\s]+', function_name):
            if len(part) >= 2:
                keywords.add(part)
        # プロセス手順名の先頭部分
        for proc in sections.procedures[:5]:
            # "(2EL)" を除去した部分
            clean = _SI_REF_PAREN.sub('', proc).strip()
            if len(clean) >= 2:
                keywords.add(clean[:20])
        return sorted(keywords)[:10]

    def _extract_en_purpose(self, en_docx_path: Path) -> str:
        """EN版docxからPurposeセクションのテキストを抽出。"""
        doc = Document(str(en_docx_path))
        in_purpose = False
        buffer = []
        for para in doc.paragraphs:
            if para.style.name == "Heading 1":
                if "Purpose" in para.text:
                    in_purpose = True
                elif in_purpose:
                    break
            elif in_purpose and para.text.strip():
                buffer.append(para.text.strip())
        return " ".join(buffer)

    async def _generate_description(
        self, purpose: str, procedures: list[str], module_context: str
    ) -> str:
        """LLM(light)でdescription生成。"""
        procs_text = "\n".join(f"- {p}" for p in procedures[:8])
        messages = [
            SystemMessage(content="あなたはSAP S/4 HANAの業務コンサルタントです。"),
            HumanMessage(content=(
                f"以下のScope Item BPDの内容から、この機能の概要を日本語で100-200文字で要約してください。\n\n"
                f"## 目的\n{purpose}\n\n"
                f"## 主要プロセスステップ\n{procs_text}\n\n"
                f"## モジュール: {module_context}\n\n"
                f"要約は「この機能は〜」で始め、主要な業務プロセスと利用シーンを含めてください。"
            )),
        ]
        result = await self._llm.call_light(messages)
        return result.content.strip()


# =============================================================================
# Module Overview PDF Parser (TASK-B03)
# =============================================================================

# Discovery WS PDFファイル名からモジュール情報を抽出するパターン
_PDF_MODULE_PATTERNS = {
    "SD": (re.compile(r'SD[ソリューション]|販売', re.IGNORECASE), "SD", "販売管理"),
    "MM": (re.compile(r'購買|調達', re.IGNORECASE), "MM", "購買管理"),
    "PP": (re.compile(r'生産|製造', re.IGNORECASE), "PP", "生産管理"),
    "PS": (re.compile(r'プロフェッショナル|サービス', re.IGNORECASE), "PS", "プロフェッショナルサービス"),
    "FI": (re.compile(r'財務|会計', re.IGNORECASE), "FI", "財務会計"),
    "EWM": (re.compile(r'倉庫|在庫管理', re.IGNORECASE), "EWM", "倉庫管理"),
}

# PDF内のScope Item ID参照パターン
_PDF_SI_PATTERN = re.compile(r'\b([A-Z0-9]{2,5})\b')


class ModuleOverviewParser:
    """Discovery WS Module Overview PDFからModuleOverviewDataを生成。"""

    def __init__(self, llm_client: LLMClient | None = None) -> None:
        self._llm = llm_client

    async def parse_module_overview(self, pdf_path: Path) -> ModuleOverviewData:
        """PDFからModuleOverview構造化データを生成。"""
        # 1. テキスト抽出
        full_text, page_count = self._extract_text(pdf_path)

        # 2. モジュール情報検出
        module, module_name = self._detect_module(pdf_path.name)

        # 3. Scope Item ID参照検出
        covers = self._detect_scope_item_references(full_text)

        # 4. サマリー生成
        summary = full_text[:500]
        if self._llm:
            summary = await self._generate_summary(full_text, module_name)

        # 5. ID生成
        id_slug = re.sub(r'[^a-z0-9]+', '-', module_name.lower()).strip('-')
        doc_id = f"MO-{module}-{id_slug}"[:50]

        return ModuleOverviewData(
            id=doc_id,
            product="SAP S/4 HANA Public Edition",
            product_namespace="SAP",
            module=module,
            module_name=module_name,
            summary=summary,
            source_doc=pdf_path.name,
            page_count=page_count,
            covers_scope_items=[f"SAP-{sid}" for sid in covers],
        )

    def _extract_text(self, pdf_path: Path) -> tuple[str, int]:
        """pdfplumberでテキスト+テーブルを抽出。"""
        full_text = ""
        page_count = 0
        with pdfplumber.open(str(pdf_path)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text() or ""
                full_text += text + "\n"
                # テーブルもテキスト化
                for table in page.extract_tables():
                    for row in table:
                        cells = [str(c) for c in row if c]
                        if cells:
                            full_text += " | ".join(cells) + "\n"
        return full_text, page_count

    def _detect_module(self, filename: str) -> tuple[str, str]:
        """ファイル名からモジュール情報を検出。"""
        for _, (pattern, module, name) in _PDF_MODULE_PATTERNS.items():
            if pattern.search(filename):
                # ファイル名から具体的なソリューション名を抽出
                # "SDソリューション(在庫販売)" → "SD在庫販売ソリューション"
                match = re.search(r'_([^_]+)\.pdf$', filename)
                specific_name = match.group(1) if match else name
                return module, specific_name
        return "OTHER", filename

    def _detect_scope_item_references(self, text: str) -> list[str]:
        """PDFテキストからScope Item IDパターンを検出。

        Discovery WS PDFでは "BD9 在庫からの販売" のようなパターンで
        Scope Item IDが記載されている。
        """
        # 2-4文字の英数字 + 直後に日本語テキスト
        pattern = re.compile(r'\b([A-Z][A-Z0-9]{1,3})\s+[\u3000-\u9fff\u30A0-\u30FF]')
        candidates = pattern.findall(text)
        # 一般的な略語を除外
        exclude = {"SAP", "ERP", "PDF", "CRM", "BOM", "API", "CSV", "URL", "GUI",
                   "RFC", "EDI", "MRP", "BTP", "USD", "EUR", "JPY", "STO", "MDS"}
        result = sorted(set(c for c in candidates if c not in exclude))
        return result

    async def _generate_summary(self, full_text: str, module_name: str) -> str:
        """LLM(light)でモジュール概要サマリーを生成。"""
        # テキスト truncation: 先頭6000 + 末尾2000
        if len(full_text) > 8000:
            excerpt = full_text[:6000] + "\n\n...(中略)...\n\n" + full_text[-2000:]
        else:
            excerpt = full_text

        messages = [
            SystemMessage(content="あなたはSAP S/4 HANAの業務コンサルタントです。"),
            HumanMessage(content=(
                f"以下のSAP S/4 HANAモジュール紹介資料から、モジュールの機能概要を日本語で300-500文字で要約してください。\n\n"
                f"## モジュール名: {module_name}\n\n"
                f"## 資料テキスト（抜粋）\n{excerpt}\n\n"
                f"要約は以下を含めてください:\n"
                f"- このモジュールがカバーする主要な業務プロセス\n"
                f"- 含まれるScope Itemの概要\n"
                f"- 他モジュールとの連携ポイント"
            )),
        ]
        result = await self._llm.call_light(messages)
        return result.content.strip()
