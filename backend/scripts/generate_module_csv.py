#!/usr/bin/env python3
"""モジュール分類マスターCSV自動生成スクリプト

全BPDのタイトルテーブルから function_name を抽出し、
キーワードベースのヒューリスティックでSAPモジュールを推定する。

Usage:
    python scripts/generate_module_csv.py [--bpd-dir DIR] [--output FILE]
"""

import argparse
import csv
import re
from pathlib import Path

from docx import Document

# SAPモジュール分類ルール（キーワード → モジュール）
# 優先度順: 先にマッチしたルールが適用される
MODULE_RULES: list[tuple[list[str], str, str, str]] = [
    # (keywords, module, module_name_ja, business_domain)
    # === SD: 販売管理 ===
    (["販売", "受注", "出荷", "請求", "得意先", "返品", "価格設定", "見積", "与信", "出荷通知"],
     "SD", "販売管理", "販売"),
    # === MM: 購買管理 ===
    (["購買", "発注", "仕入先", "入庫", "調達", "外注", "購入", "サプライヤ"],
     "MM", "購買管理", "購買"),
    # === PP: 生産管理 ===
    (["生産", "製造", "MRP", "BOM", "作業手順", "計画手配", "組立", "原料", "製品マスタ"],
     "PP", "生産管理", "生産"),
    # === FI: 財務会計 ===
    (["会計", "仕訳", "元帳", "勘定", "決算", "消費税", "税務", "請求書照合", "支払", "債権", "債務",
      "会社間", "連結", "固定資産"],
     "FI", "財務会計", "財務"),
    # === CO: 管理会計 ===
    (["原価", "利益", "コストセンタ", "内部指図", "活動配賦", "配賦", "収益性"],
     "CO", "管理会計", "管理会計"),
    # === WM/EWM: 倉庫管理 ===
    (["倉庫", "在庫", "棚卸", "保管場所", "入出庫", "ピッキング", "物理在庫"],
     "EWM", "倉庫管理", "物流"),
    # === TM: 輸送管理 ===
    (["輸送", "配送", "出荷ポイント", "運送"],
     "TM", "輸送管理", "物流"),
    # === PM: プラント保全 ===
    (["保全", "メンテナンス", "設備", "点検", "修理", "故障"],
     "PM", "プラント保全", "保全"),
    # === QM: 品質管理 ===
    (["品質", "検査", "品質管理", "不良", "ロット"],
     "QM", "品質管理", "品質"),
    # === PS: プロジェクト管理 ===
    (["プロジェクト", "WBS", "マイルストーン", "サービス"],
     "PS", "プロジェクト管理", "プロジェクト"),
    # === HR/HCM: 人事管理 ===
    (["人事", "給与", "勤怠", "従業員", "採用", "タレント"],
     "HCM", "人事管理", "人事"),
    # === TR: 資金管理 ===
    (["資金", "キャッシュ", "為替", "銀行", "資金管理", "ヘッジ", "デリバティブ", "外国為替"],
     "TR", "資金管理", "財務"),
    # === GRC: ガバナンス・リスク ===
    (["コンプライアンス", "監査", "リスク", "ガバナンス", "権限"],
     "GRC", "ガバナンス", "ガバナンス"),
    # === BASIS: 基盤 ===
    (["設定", "カスタマイズ", "移行", "データ移行", "移送"],
     "BASIS", "システム基盤", "基盤"),
    # === Analytics ===
    (["分析", "レポート", "ダッシュボード", "KPI", "アナリティクス"],
     "ANA", "分析", "分析"),
]


def extract_function_name(docx_path: Path) -> str | None:
    """BPD docxのタイトルテーブル（Table 0）からfunction_nameを抽出。"""
    try:
        doc = Document(str(docx_path))
        if not doc.tables:
            return None
        # Table 0 の2行目1列目にタイトルがある
        table = doc.tables[0]
        if len(table.rows) >= 2:
            title = table.rows[1].cells[-1].text.strip()
            # "(XXX_JP)" サフィックスを除去
            title = re.sub(r'\s*\([A-Z0-9]+_JP\)\s*$', '', title)
            return title
        return None
    except Exception:
        return None


def classify_module(function_name: str) -> tuple[str, str, str]:
    """function_name からSAPモジュールを推定。"""
    for keywords, module, module_name_ja, business_domain in MODULE_RULES:
        for kw in keywords:
            if kw in function_name:
                return module, module_name_ja, business_domain
    return "OTHER", "その他", "その他"


def discover_bpd_prefixes(bpd_dir: Path) -> dict[str, Path]:
    """BPDディレクトリからJA版docxのプレフィクス→パスを取得。"""
    result = {}
    for f in sorted(bpd_dir.iterdir()):
        if f.is_file() and "_BPD_JA_" in f.name and f.name.endswith(".docx"):
            prefix = f.name.split("_S4CLD")[0]
            result[prefix] = f
    return result


def main():
    parser = argparse.ArgumentParser(description="Generate module classification CSV")
    parser.add_argument(
        "--bpd-dir",
        type=Path,
        default=Path(__file__).resolve().parent.parent.parent / "product_doc",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path(__file__).resolve().parent.parent / "data" / "module_classification_sap.csv",
    )
    args = parser.parse_args()

    prefixes = discover_bpd_prefixes(args.bpd_dir)
    print(f"Found {len(prefixes)} BPD JA docx files")

    records = []
    module_stats = {}
    errors = 0

    for prefix, docx_path in prefixes.items():
        fn = extract_function_name(docx_path)
        if fn is None:
            print(f"  WARN: {prefix} - could not extract function_name")
            errors += 1
            module, module_ja, domain = "OTHER", "その他", "その他"
        else:
            module, module_ja, domain = classify_module(fn)

        records.append({
            "scope_item_prefix": prefix,
            "module": module,
            "module_name_ja": module_ja,
            "business_domain": domain,
            "product": "SAP",
        })
        module_stats[module] = module_stats.get(module, 0) + 1

    # CSV出力
    args.output.parent.mkdir(parents=True, exist_ok=True)
    with open(args.output, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=["scope_item_prefix", "module", "module_name_ja", "business_domain", "product"],
        )
        writer.writeheader()
        writer.writerows(records)

    print(f"\nCSV written: {args.output}")
    print(f"Total: {len(records)} records ({errors} extraction errors)")
    print(f"\nModule distribution:")
    for module, count in sorted(module_stats.items(), key=lambda x: -x[1]):
        print(f"  {module:8s}: {count}")


if __name__ == "__main__":
    main()
