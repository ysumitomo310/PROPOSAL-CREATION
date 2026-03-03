#!/usr/bin/env python3
"""BPDフォーマットSpike（TASK-B02a）

BPDドキュメントの構造を解析し、design.md の前提と一致するか検証する。
結果をコンソールに出力する。

Usage:
    python scripts/spike_bpd_structure.py [--bpd-dir DIR] [--samples N]
"""

import argparse
import json
import random
import re
from collections import Counter
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


def inspect_docx(docx_path: Path) -> dict:
    """docxファイルの構造を解析する。"""
    doc = Document(str(docx_path))

    # 1. パラグラフ: スタイル/テキスト
    paragraphs = []
    style_counter = Counter()
    for para in doc.paragraphs:
        style = para.style.name
        style_counter[style] += 1
        if style.startswith("Heading") or para.text.strip():
            paragraphs.append({
                "style": style,
                "text": para.text[:120],
            })

    # 2. 見出し階層
    headings = [
        {"level": p["style"], "text": p["text"]}
        for p in paragraphs
        if p["style"].startswith("Heading")
    ]

    # 3. セクション名検索（design.md 前提）
    section_keywords = {
        "Purpose": [],
        "Prerequisites": [],
        "Business Conditions": [],
        "Procedure": [],
        "目的": [],
        "前提条件": [],
        "業務条件": [],
        "手順": [],
        "Scope": [],
    }
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\n"
        for kw in section_keywords:
            if kw.lower() in para.text.lower():
                section_keywords[kw].append(para.text[:100])

    # 4. テーブル情報
    tables = []
    for i, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns)
        headers = []
        if rows > 0:
            headers = [cell.text.strip()[:50] for cell in table.rows[0].cells]
        # サンプル行
        sample_row = []
        if rows > 1:
            sample_row = [cell.text.strip()[:50] for cell in table.rows[1].cells]
        tables.append({
            "index": i,
            "rows": rows,
            "cols": cols,
            "headers": headers,
            "sample_row": sample_row,
        })

    # 5. Scope Item ID パターン検出
    # 2-4文字の英数字パターン（BPD内の他SI参照）
    si_pattern = re.compile(r'\b([A-Z0-9]{2,4})\b')
    si_candidates = si_pattern.findall(full_text)
    si_counter = Counter(si_candidates)
    # 出現頻度2回以上のものだけ
    frequent_ids = {k: v for k, v in si_counter.items() if v >= 2 and len(k) >= 2}

    # 6. モジュール名検出（SD, MM, FI, CO, PP, PM, QM, PS, WM, HR等）
    module_pattern = re.compile(
        r'\b(SD|MM|FI|CO|PP|PM|QM|PS|WM|HR|EWM|TM|LE|CS|RE|TR|GRC|BW|IS-'
        r'|Sales|Procurement|Finance|Manufacturing|Warehouse|Logistics'
        r'|販売|購買|会計|生産|在庫|物流|原価|人事|品質)\b',
        re.IGNORECASE,
    )
    module_mentions = module_pattern.findall(full_text)
    module_counter = Counter(module_mentions)

    return {
        "file": docx_path.name,
        "paragraph_count": len(doc.paragraphs),
        "style_distribution": dict(style_counter.most_common(15)),
        "headings": headings,
        "section_keywords": {k: v for k, v in section_keywords.items() if v},
        "tables": tables,
        "frequent_scope_item_ids": dict(sorted(frequent_ids.items(), key=lambda x: -x[1])[:20]),
        "module_mentions": dict(module_counter.most_common(10)),
        "text_length": len(full_text),
    }


def inspect_xlsx(xlsx_path: Path) -> dict:
    """xlsxファイルの構造を解析する。"""
    wb = load_workbook(str(xlsx_path), data_only=True)
    sheets = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(max_row=5, values_only=True))
        sheets.append({
            "name": sheet_name,
            "dimensions": ws.dimensions,
            "sample_rows": [
                [str(cell)[:40] if cell is not None else "" for cell in row]
                for row in rows
            ],
        })
    wb.close()
    return {"file": xlsx_path.name, "sheets": sheets}


def discover_bpd_sets(bpd_dir: Path) -> list[dict]:
    """BPD 3ファイルセットを検出する。"""
    ja_files = {}
    en_files = {}
    xlsx_files = {}

    for f in bpd_dir.iterdir():
        if not f.is_file():
            continue
        name = f.name
        if "_BPD_JA_" in name and name.endswith(".docx"):
            prefix = name.split("_S4CLD")[0]
            ja_files[prefix] = f
        elif "_BPD_EN_" in name and name.endswith(".docx"):
            prefix = name.split("_S4CLD")[0]
            en_files[prefix] = f
        elif "_BPD_EN_" in name and name.endswith(".xlsx"):
            prefix = name.split("_S4CLD")[0]
            xlsx_files[prefix] = f

    sets = []
    for prefix in sorted(ja_files.keys()):
        entry = {"prefix": prefix, "ja": ja_files[prefix]}
        if prefix in en_files:
            entry["en"] = en_files[prefix]
        if prefix in xlsx_files:
            entry["xlsx"] = xlsx_files[prefix]
        sets.append(entry)

    return sets


def main():
    parser = argparse.ArgumentParser(description="BPD Structure Spike")
    parser.add_argument(
        "--bpd-dir",
        type=Path,
        default=Path(__file__).resolve().parent.parent.parent / "product_doc",
    )
    parser.add_argument("--samples", type=int, default=5, help="ランダムサンプル数")
    args = parser.parse_args()

    bpd_dir = args.bpd_dir
    print(f"=== BPD Structure Spike ===")
    print(f"BPD directory: {bpd_dir}")

    # ファイルセット発見
    sets = discover_bpd_sets(bpd_dir)
    print(f"\nTotal BPD sets found: {len(sets)}")

    # ファイルセット構成統計
    complete = sum(1 for s in sets if "en" in s and "xlsx" in s)
    ja_only = sum(1 for s in sets if "en" not in s)
    no_xlsx = sum(1 for s in sets if "xlsx" not in s and "en" in s)
    print(f"  Complete (JA+EN+xlsx): {complete}")
    print(f"  JA only (no EN): {ja_only}")
    print(f"  No xlsx: {no_xlsx}")

    # サンプリング
    if len(sets) > args.samples:
        sample_sets = random.sample(sets, args.samples)
    else:
        sample_sets = sets

    print(f"\n--- Analyzing {len(sample_sets)} samples ---\n")

    all_section_keys = Counter()
    all_module_mentions = Counter()
    all_heading_patterns = []

    for i, bpd_set in enumerate(sample_sets):
        print(f"\n{'='*60}")
        print(f"[{i+1}/{len(sample_sets)}] Prefix: {bpd_set['prefix']}")
        print(f"{'='*60}")

        # JA docx
        print(f"\n--- JA DOCX: {bpd_set['ja'].name} ---")
        ja_result = inspect_docx(bpd_set["ja"])
        print(f"  Paragraphs: {ja_result['paragraph_count']}")
        print(f"  Text length: {ja_result['text_length']} chars")
        print(f"  Styles: {json.dumps(ja_result['style_distribution'], ensure_ascii=False, indent=4)}")
        print(f"  Headings ({len(ja_result['headings'])}):")
        for h in ja_result["headings"]:
            print(f"    [{h['level']}] {h['text']}")
        print(f"  Section keywords found: {json.dumps(ja_result['section_keywords'], ensure_ascii=False, indent=4)}")
        print(f"  Tables ({len(ja_result['tables'])}):")
        for t in ja_result["tables"]:
            print(f"    Table {t['index']}: {t['rows']}rows x {t['cols']}cols | headers={t['headers'][:5]}")
            if t["sample_row"]:
                print(f"      sample: {t['sample_row'][:5]}")
        print(f"  Module mentions: {json.dumps(ja_result['module_mentions'], ensure_ascii=False)}")
        print(f"  Frequent SI IDs: {json.dumps(dict(list(ja_result['frequent_scope_item_ids'].items())[:10]), ensure_ascii=False)}")

        for k in ja_result["section_keywords"]:
            all_section_keys[k] += 1
        for k, v in ja_result["module_mentions"].items():
            all_module_mentions[k] += v
        all_heading_patterns.append([h["level"] for h in ja_result["headings"]])

        # EN docx
        if "en" in bpd_set:
            print(f"\n--- EN DOCX: {bpd_set['en'].name} ---")
            en_result = inspect_docx(bpd_set["en"])
            print(f"  Paragraphs: {en_result['paragraph_count']}")
            print(f"  Headings ({len(en_result['headings'])}):")
            for h in en_result["headings"][:10]:
                print(f"    [{h['level']}] {h['text']}")
            print(f"  Section keywords found: {list(en_result['section_keywords'].keys())}")

        # xlsx
        if "xlsx" in bpd_set:
            print(f"\n--- XLSX: {bpd_set['xlsx'].name} ---")
            xlsx_result = inspect_xlsx(bpd_set["xlsx"])
            for sheet in xlsx_result["sheets"]:
                print(f"  Sheet: {sheet['name']} | Dims: {sheet['dimensions']}")
                for j, row in enumerate(sheet["sample_rows"]):
                    print(f"    Row {j}: {row[:6]}")

    # サマリー
    print(f"\n{'='*60}")
    print(f"=== SUMMARY ===")
    print(f"{'='*60}")
    print(f"\nSection keywords found across {len(sample_sets)} samples:")
    for k, v in all_section_keys.most_common():
        print(f"  {k}: {v}/{len(sample_sets)} samples")
    print(f"\nModule mentions (aggregated):")
    for k, v in all_module_mentions.most_common(15):
        print(f"  {k}: {v}")
    print(f"\nHeading patterns:")
    for i, pattern in enumerate(all_heading_patterns):
        print(f"  Sample {i+1}: {pattern[:10]}{'...' if len(pattern) > 10 else ''}")


if __name__ == "__main__":
    main()
